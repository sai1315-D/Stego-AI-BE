import unittest
import io
import os
import tempfile
import numpy as np
import cv2
import wave
import docx
from pypdf import PdfWriter

from app.services.stego_image import ImageStegoDetector
from app.services.stego_audio import AudioStegoDetector
from app.services.stego_video import VideoStegoDetector
from app.services.stego_document import DocumentStegoDetector
from app.services.ai_engine import ai_engine

class TestStegoDetectors(unittest.TestCase):

    def setUp(self):
        # Ensure model is initialized
        self.assertIsNotNone(ai_engine)

    def test_image_stego_detector(self):
        # 1. Create a clean 100x100 image
        img = np.zeros((100, 100, 3), dtype=np.uint8)
        # Add a gradient to make it look like a natural image with some correlation
        for y in range(100):
            for x in range(100):
                img[y, x] = [x * 2, y * 2, (x + y) % 256]
                
        _, png_bytes = cv2.imencode('.png', img)
        res = ImageStegoDetector.analyze(png_bytes.tobytes())
        
        self.assertEqual(res["file_type"], "image")
        self.assertIn("lsb_entropy", res["metrics"])
        self.assertIn("chi_square_stat", res["metrics"])
        
        # Test engine integration
        engine_res = ai_engine.analyze_file(png_bytes.tobytes(), "test_img.png", "image")
        self.assertIn("risk_level", engine_res)
        self.assertIn("risk_score", engine_res)

    def test_audio_stego_detector(self):
        # 2. Create a clean 1-second sine wave audio in WAV format
        audio_io = io.BytesIO()
        with wave.open(audio_io, 'wb') as wav:
            wav.setnchannels(1)
            wav.setsampwidth(2)
            wav.setframerate(22050)
            
            # 1 second of 440Hz sine wave
            t = np.linspace(0, 1, 22050, endpoint=False)
            samples = np.int16(np.sin(2 * np.pi * 440.0 * t) * 16384)
            wav.writeframes(samples.tobytes())
            
        audio_bytes = audio_io.getvalue()
        res = AudioStegoDetector.analyze(audio_bytes)
        
        self.assertEqual(res["file_type"], "audio")
        self.assertIn("audio_lsb_correlation", res["metrics"])
        self.assertIn("mfcc_ratio", res["metrics"])
        
        # Test engine integration
        engine_res = ai_engine.analyze_file(audio_bytes, "test_audio.wav", "audio")
        self.assertIn("risk_level", engine_res)

    def test_video_stego_detector(self):
        # 3. Create a clean dummy video file with 5 frames
        temp_dir = tempfile.gettempdir()
        temp_video_path = os.path.join(temp_dir, "test_temp_video.mp4")
        
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(temp_video_path, fourcc, 10.0, (60, 60))
        
        for i in range(5):
            # Create a frames with slight changes to prevent zero variance errors
            frame = np.zeros((60, 60, 3), dtype=np.uint8)
            frame[:, :] = [i * 10, i * 20, i * 30]
            out.write(frame)
        out.release()
        
        with open(temp_video_path, "rb") as f:
            video_bytes = f.read()
            
        if os.path.exists(temp_video_path):
            os.remove(temp_video_path)
            
        res = VideoStegoDetector.analyze(video_bytes)
        self.assertEqual(res["file_type"], "video")
        self.assertIn("entropy_variance", res["metrics"])
        
        # Test engine integration
        engine_res = ai_engine.analyze_file(video_bytes, "test_video.mp4", "video")
        self.assertIn("risk_level", engine_res)

    def test_document_stego_detector_txt(self):
        # 4a. Clean plain text
        clean_text = "This is clean text containing no hidden stego channel."
        res_clean = DocumentStegoDetector.analyze(clean_text.encode("utf-8"), "test.txt")
        self.assertEqual(res_clean["risk_level"], "SAFE")
        self.assertEqual(res_clean["metrics"]["zero_width_chars"], 0)
        
        # 4b. Text containing multiple zero-width characters and whitespace encoded lines (Stego)
        stego_text = (
            "This is stego text containing hidden payload \u200b\u200b\u200b\u200b\u200b\u200b\u200b\u200b characters.\n"
            "Line 1 with steganographic trailing spaces    \n"
            "Line 2 with steganographic trailing spaces    \n"
            "Line 3 with steganographic trailing spaces    \n"
        )
        res_stego = DocumentStegoDetector.analyze(stego_text.encode("utf-8"), "test.txt")
        self.assertEqual(res_stego["risk_level"], "DANGEROUS")
        self.assertGreater(res_stego["metrics"]["zero_width_chars"], 0)

    def test_document_stego_detector_docx(self):
        # 5. Create a word document with multiple hidden text runs and zero-width characters to exceed DANGEROUS threshold
        doc = docx.Document()
        p = doc.add_paragraph("This is the main visible text in paragraph with hidden chars \u200b\u200b\u200b\u200b\u200b\u200b\u200b\u200b\u200b\u200b. ")
        for i in range(5):
            p = doc.add_paragraph()
            run = p.add_run(f"Hidden run {i}")
            run.font.hidden = True
        
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_bytes = docx_io.getvalue()
        
        res = DocumentStegoDetector.analyze(docx_bytes, "test.docx")
        self.assertEqual(res["file_type"], "document")
        self.assertGreater(res["metrics"]["hidden_runs"], 0)
        self.assertEqual(res["risk_level"], "DANGEROUS")

    def test_document_stego_detector_pdf(self):
        # 6. Create a PDF and append data at the end (after %%EOF)
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        
        pdf_io = io.BytesIO()
        writer.write(pdf_io)
        pdf_clean_bytes = pdf_io.getvalue()
        
        res_clean = DocumentStegoDetector.analyze(pdf_clean_bytes, "test.pdf")
        self.assertEqual(res_clean["metrics"]["appended_bytes"], 0)
        
        # Stego PDF (appended extra payload bytes after EOF marker)
        pdf_stego_bytes = pdf_clean_bytes + b"\n%%EOF\nHidden payload content injected in tail bytes."
        res_stego = DocumentStegoDetector.analyze(pdf_stego_bytes, "test.pdf")
        self.assertGreater(res_stego["metrics"]["appended_bytes"], 0)
        # 46 appended bytes falls under SUSPICIOUS risk classification
        self.assertIn(res_stego["risk_level"], ["SUSPICIOUS", "DANGEROUS"])

if __name__ == "__main__":
    unittest.main()
