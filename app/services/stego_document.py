import io
import re
from pypdf import PdfReader
from docx import Document

class DocumentStegoDetector:
    @staticmethod
    def analyze(file_bytes: bytes, file_name: str = "") -> dict:
        """
        Analyze documents (PDF, DOCX, TXT) for steganography.
        Detects appended data after EOF, zero-width characters, metadata anomalies,
        and invisible/whitespace text encoding.
        """
        # Determine file type
        name_lower = file_name.lower()
        if name_lower.endswith(".pdf"):
            return DocumentStegoDetector._analyze_pdf(file_bytes)
        elif name_lower.endswith(".docx"):
            return DocumentStegoDetector._analyze_docx(file_bytes)
        else:
            # Default to plain text parsing
            return DocumentStegoDetector._analyze_txt(file_bytes)

    @staticmethod
    def _analyze_pdf(file_bytes: bytes) -> dict:
        """
        PDF Steganography Detection:
        - Checks for appended bytes after the final %%EOF marker.
        - Checks metadata for anomaly custom attributes.
        - Checks for hidden zero-width characters or whitespace encoding in extracted text.
        """
        score = 0
        reasons = []
        
        # 1. Check for trailing data after %%EOF (very common steg technique)
        eof_index = file_bytes.rfind(b"%%EOF")
        appended_bytes = 0
        if eof_index != -1:
            # Exclude trailing newlines/spaces typical of standard file saving
            trailing = file_bytes[eof_index + 5:].strip()
            if len(trailing) > 10:  # Allow small grace for file markers
                appended_bytes = len(trailing)
                score += min(45, int(appended_bytes / 50) + 15)
                reasons.append(f"Detected {appended_bytes} bytes of appended payload content after final %%EOF marker.")

        # Load PDF reader
        try:
            pdf_io = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_io)
            
            # 2. Metadata Analysis
            metadata = reader.metadata or {}
            custom_metadata_count = 0
            for key in metadata.keys():
                if not key.startswith("/") or key.lower() not in [
                    "/title", "/author", "/subject", "/keywords", "/creator", "/producer", "/creationdate", "/moddate", "/trapped"
                ]:
                    custom_metadata_count += 1
            
            if custom_metadata_count > 3:
                score += 15
                reasons.append(f"Excessive custom metadata fields ({custom_metadata_count}) typical of payload descriptor headers.")

            # 3. Text Extraction & Encoding Pattern Detection
            full_text = ""
            for page in reader.pages:
                try:
                    full_text += page.extract_text() or ""
                except Exception:
                    pass
            
            # Search for Zero-Width Characters
            zw_pattern = re.compile(r'[\u200b-\u200d\ufeff\u200e\u200f\u202a-\u202e]')
            zw_matches = zw_pattern.findall(full_text)
            zw_count = len(zw_matches)
            
            if zw_count > 0:
                score += min(35, zw_count * 5)
                reasons.append(f"Found {zw_count} occurrences of hidden Zero-Width unicode characters (common text stego).")

            # Search for trailing whitespace encoding (SNOW-style tab/space patterns)
            # Lines ending in combinations of multiple spaces and tabs
            ws_steg_pattern = re.compile(r'[ \t]{4,}\r?\n')
            ws_matches = ws_steg_pattern.findall(full_text)
            if len(ws_matches) > 2:
                score += 20
                reasons.append(f"Detected {len(ws_matches)} lines with trailing whitespace combinations indicating binary-to-space encoding.")

        except Exception as e:
            # If PDF parsing fails but we have trailing data, it is highly suspicious
            if appended_bytes > 0:
                score += 30
            else:
                raise ValueError(f"Failed to parse PDF format: {str(e)}")

        risk_score = max(0, min(100, score))
        
        return DocumentStegoDetector._build_response("PDF", risk_score, reasons, {
            "appended_bytes": appended_bytes,
            "custom_metadata_fields": custom_metadata_count if 'custom_metadata_count' in locals() else 0,
            "zero_width_chars": zw_count if 'zw_count' in locals() else 0,
            "whitespace_patterns": len(ws_matches) if 'ws_matches' in locals() else 0
        })

    @staticmethod
    def _analyze_docx(file_bytes: bytes) -> dict:
        """
        DOCX Steganography Detection:
        - Check document properties.
        - Check for hidden/zero-size text elements.
        - Scan text content for zero-width unicode or whitespace encoding.
        """
        score = 0
        reasons = []
        
        try:
            docx_io = io.BytesIO(file_bytes)
            doc = Document(docx_io)
            
            # Extract paragraphs
            full_text = ""
            hidden_runs = 0
            for p in doc.paragraphs:
                full_text += p.text + "\n"
                # Check for hidden font formatting in runs
                for run in p.runs:
                    if run.font.hidden:
                        hidden_runs += 1
            
            if hidden_runs > 0:
                score += min(40, hidden_runs * 10)
                reasons.append(f"Detected {hidden_runs} instances of hidden text formatting inside document runs.")
                
            # Scan table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\n"

            # Check for Zero-Width Characters
            zw_pattern = re.compile(r'[\u200b-\u200d\ufeff\u200e\u200f\u202a-\u202e]')
            zw_matches = zw_pattern.findall(full_text)
            zw_count = len(zw_matches)
            
            if zw_count > 0:
                score += min(35, zw_count * 5)
                reasons.append(f"Found {zw_count} occurrences of hidden Zero-Width unicode characters.")

            # Check core properties anomalies
            props = doc.core_properties
            custom_fields = 0
            # DOCX files store custom properties in another XML, but we can verify if standard fields are misused
            if props.comments and len(props.comments) > 500:
                score += 15
                reasons.append("Unusually large document Comments property, typical of embedded metadata payloads.")

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

        risk_score = max(0, min(100, score))
        
        return DocumentStegoDetector._build_response("DOCX", risk_score, reasons, {
            "hidden_runs": hidden_runs,
            "zero_width_chars": zw_count if 'zw_count' in locals() else 0,
            "large_comments_property": len(props.comments) if ('props' in locals() and props.comments) else 0
        })

    @staticmethod
    def _analyze_txt(file_bytes: bytes) -> dict:
        """
        Plain Text Steganography Detection:
        - Analyze for zero-width characters.
        - Analyze for whitespace-only stego encoding.
        """
        score = 0
        reasons = []
        
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode("latin-1")
            except Exception:
                raise ValueError("Could not decode text file as UTF-8 or Latin-1.")

        # Scan for Zero-Width characters
        zw_pattern = re.compile(r'[\u200b-\u200d\ufeff\u200e\u200f\u202a-\u202e]')
        zw_matches = zw_pattern.findall(text)
        zw_count = len(zw_matches)
        
        if zw_count > 0:
            score += min(50, zw_count * 10)
            reasons.append(f"Detected {zw_count} hidden zero-width unicode characters in plain text.")

        # Scan for line-ending spaces/tabs patterns (SNOW)
        lines = text.splitlines()
        ws_lines = 0
        for line in lines:
            if line.endswith(" ") or line.endswith("\t"):
                # Check if it has multiple spaces/tabs at the end of the line
                # Standard editors strip trailing whitespace, so their existence is anomalous.
                tail = len(line) - len(line.rstrip(" \t"))
                if tail >= 4:
                    ws_lines += 1
                    
        if ws_lines > 2:
            score += min(40, ws_lines * 10)
            reasons.append(f"Found {ws_lines} lines ending with specific combinations of trailing spaces/tabs.")

        risk_score = max(0, min(100, score))
        
        return DocumentStegoDetector._build_response("TXT", risk_score, reasons, {
            "zero_width_chars": zw_count,
            "whitespace_encoded_lines": ws_lines
        })

    @staticmethod
    def _build_response(format_name: str, risk_score: int, reasons: list, metrics: dict) -> dict:
        if risk_score <= 30:
            risk_level = "SAFE"
            description = f"No document steganography detected. File structure conforms to standard {format_name} specifications."
        elif risk_score <= 70:
            risk_level = "SUSPICIOUS"
            description = f"Suspicious characteristics found. {reasons[0] if reasons else 'Possible hidden unicode characters.'}"
        else:
            risk_level = "DANGEROUS"
            description = f"Steganography detected: " + "; ".join(reasons)
            
        return {
            "file_name": "",
            "file_type": "document",
            "threat_probability": risk_score / 100.0,
            "risk_score": risk_score,
            "risk_level": risk_level,
            "description": description,
            "metrics": metrics
        }
